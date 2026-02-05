import os
import streamlit as st
from openai import OpenAI
from docx import Document
from io import BytesIO
import fitz  # PyMuPDF
import time

# --- 1. НАСТРОЙКА СТРАНИЦЫ ---
st.set_page_config(
    page_title="Симулятор Вредного Заказчика", 
    layout="wide", 
    page_icon="🧛"
)

# --- 2. ФУНКЦИИ ОБРАБОТКИ ТЕКСТА ---
def extract_text_from_pdf(file):
    """Извлекает текст из PDF с пометкой страниц."""
    doc = fitz.open(stream=file.read(), filetype="pdf")
    text = ""
    for i, page in enumerate(doc):
        text += f"\n[СТРАНИЦА {i+1}]\n{page.get_text()}"
    return text

def extract_text_from_docx(file):
    """Извлекает текст из DOCX с пометкой абзацев."""
    doc = Document(file)
    text = ""
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            text += f"[Абзац {i+1}] {para.text}\n"
    return text

def load_bad_history():
    """Загружает базу типичных отказов из текстового файла."""
    if os.path.exists("bad_history.txt"):
        try:
            with open("bad_history.txt", "r", encoding="utf-8") as f:
                return f.read()
        except:
            return "База прошлых отказов пуста."
    return "База прошлых отказов пуста."

def create_docx(text):
    """Создает Word-файл из текста отчета."""
    doc = Document()
    doc.add_heading('ПРОТОКОЛ НЕСООТВЕТСТВИЙ', 0)
    # Разбиваем текст на параграфы для корректного отображения в Word
    for line in text.split('\n'):
        doc.add_paragraph(line)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- 3. ИНТЕРФЕЙС (SIDEBAR) ---
with st.sidebar:
    st.header("⚙️ Настройки AI")
    selected_model = st.selectbox(
        "Выберите модель DeepSeek:",
        ("deepseek-chat", "deepseek-reasoner"),
        index=0,
        help="chat (V3) — быстро и четко. reasoner (R1) — глубокая логика."
    )
    
    st.info("Убедитесь, что DEEPSEEK_API_KEY добавлен в Secrets в настройках Streamlit Cloud.")

# --- 4. ОСНОВНОЙ ИНТЕРФЕЙС ---
st.title("🧛 Симулятор Вредного Заказчика")
st.markdown("### Глубокий аудит соответствия Отчета требованиям Контракта")

col1, col2 = st.columns(2)
with col1:
    contract_file = st.file_uploader("📄 КОНТРАКТ (PDF/DOCX)", type=['pdf', 'docx'], key="contract")
with col2:
    report_file = st.file_uploader("📝 ОТЧЕТ ИСПОЛНИТЕЛЯ (PDF/DOCX)", type=['pdf', 'docx'], key="report")

# --- 5. ЛОГИКА АНАЛИЗА ---
if st.button("🚀 ЗАПУСТИТЬ ТОТАЛЬНЫЙ АУДИТ"):
    if contract_file and report_file:
        try:
            # Получение ключа
            api_key_val = st.secrets.get("DEEPSEEK_API_KEY")
            if not api_key_val:
                st.error("Критическая ошибка: Ключ API не найден.")
                st.stop()

            # Инициализация клиента
            client = OpenAI(base_url="https://api.deepseek.com", api_key=api_key_val)
            
            progress_bar = st.progress(0)
            status_text = st.empty()

            # Шаг 1: Извлечение текста
            status_text.info("📂 Шаг 1/4: Чтение документов...")
            c_text = extract_text_from_pdf(contract_file) if contract_file.name.endswith('.pdf') else extract_text_from_docx(contract_file)
            r_text = extract_text_from_pdf(report_file) if report_file.name.endswith('.pdf') else extract_text_from_docx(report_file)
            progress_bar.progress(25)

            # Шаг 2: Формирование промпта
            status_text.info("⚖️ Шаг 2/4: Подготовка экспертных инструкций...")
            bad_history = load_bad_history()
            
            system_prompt = f"""Ты — придирчивый инспектор госконтрактов. 
Твоя задача: найти несоответствия в Отчете относительно Контракта.

ОБЯЗАТЕЛЬНЫЙ ФОРМАТ КАЖДОГО ПУНКТА:
1. **Нарушение**: (суть ошибки)
2. **Локация**: (Страница №, Абзац № в документе)
3. **Основание**: (пункт ТЗ или пункт из Базы Отказов)
4. **Риск**: (почему это критично)

ПРОВЕРЬ ПО ТРЕМ СЛОЯМ:
- СЛОЙ 1 (База отказов): {bad_history}
- СЛОЙ 2 (Техзадание): Сверь даты, цифры, объемы, наличие всех приложений.
- СЛОЙ 3 (Формализм): ИНК, печати, стоп-слова ("около", "не менее"), полные названия (РФ -> Российская Федерация).
"""

            user_content = f"""ВНИМАНИЕ: Используй маркеры [СТРАНИЦА X] и [Абзац Y] для указания локации ошибок.

ТЕКСТ КОНТРАКТА:
{c_text[:12000]}

ТЕКСТ ОТЧЕТА ДЛЯ ПРОВЕРКИ:
{r_text[:12000]}"""

            progress_bar.progress(50)

            # Шаг 3: Запрос к API
            status_text.info(f"🧠 Шаг 3/4: Работает модель {selected_model}...")
            
            response = client.chat.completions.create(
                model=selected_model,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_content}
                ],
                max_tokens=4000,
                temperature=0.1
            )
            
            # Извлечение рассуждений (если выбрана модель reasoner)
            reasoning = getattr(response.choices[0].message, 'reasoning_content', None)
            result_text = response.choices[0].message.content
            
            progress_bar.progress(85)

            # Шаг 4: Вывод результатов
            status_text.success("✅ Аудит завершен!")
            progress_bar.progress(100)

            if reasoning:
                with st.expander("🔍 Посмотреть процесс логического размышления ИИ"):
                    st.write(reasoning)

            st.divider()
            st.subheader("📋 Протокол выявленных несоответствий")
            st.markdown(result_text)

            # Кнопка скачивания
            st.download_button(
                label="📥 Скачать протокол в Word",
                data=create_docx(result_text),
                file_name="Audit_Protocol_DeepSeek.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        except Exception as e:
            st.error(f"⚠️ Ошибка при выполнении аудита: {str(e)}")
    else:
        st.warning("⚠️ Пожалуйста, загрузите оба файла (Контракт и Отчет).")
